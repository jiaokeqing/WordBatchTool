from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ..schemas import BuiltInTemplate, TemplatePreview


def apply_builtin_template(input_path: Path, output_path: Path, template: BuiltInTemplate) -> None:
    document = Document(input_path)
    remove_empty_paragraphs(document)
    for section in document.sections:
        section.top_margin = Cm(template.margin_top_cm)
        section.bottom_margin = Cm(template.margin_bottom_cm)
        section.left_margin = Cm(template.margin_left_cm)
        section.right_margin = Cm(template.margin_right_cm)

    for index, paragraph in enumerate(document.paragraphs):
        level = detect_paragraph_level(paragraph.text, index)
        font_name, font_size = font_for_level(template, level)
        paragraph.style = document.styles["Normal"]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(template.line_spacing_pt)
        paragraph_format.space_before = Pt(template.space_before_pt)
        paragraph_format.space_after = Pt(template.space_after_pt)
        set_paragraph_line_spacing(paragraph, before_lines=0, after_lines=0)
        paragraph_format.left_indent = Pt(0)
        paragraph_format.right_indent = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
        set_first_line_indent_chars(paragraph, template.first_line_indent_chars if level == "body" else 0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == "title" else WD_ALIGN_PARAGRAPH.LEFT

        for run in paragraph.runs:
            if template.normalize_spacing or template.normalize_parentheses:
                run.text = normalize_text(run.text, template)
            apply_run_font(run, font_name, template.latin_font, font_size)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def apply_sample_template(input_path: Path, output_path: Path, sample_path: Path) -> None:
    sample = Document(sample_path)
    builtin = BuiltInTemplate()
    for paragraph in sample.paragraphs:
        if not paragraph.runs:
            continue
        font = paragraph.runs[0].font
        if font.name:
            builtin.body_font = font.name
        if font.size:
            builtin.body_size = int(font.size.pt)
        break
    apply_builtin_template(input_path, output_path, builtin)


def preview_template(sample_path: Path) -> TemplatePreview:
    document = Document(sample_path)
    styles = sorted({paragraph.style.name for paragraph in document.paragraphs if paragraph.style})
    return TemplatePreview(styles=styles[:50], paragraphs=len(document.paragraphs), tables=len(document.tables))


def remove_empty_paragraphs(document: Document) -> None:
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip():
            continue
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def detect_paragraph_level(text: str, index: int) -> str:
    stripped = text.strip()
    if not stripped:
        return "body"
    if is_salutation(stripped):
        return "salutation"
    if index == 0 and is_probable_title(stripped):
        return "title"
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        return "heading1"
    if re.match(r"^（[一二三四五六七八九十]+）", stripped) or re.match(r"^\([一二三四五六七八九十]+\)", stripped):
        return "heading2"
    if re.match(r"^\d+[\.．、]", stripped):
        return "heading3"
    return "body"


def is_salutation(text: str) -> bool:
    return text.endswith(("：", ":"))


def is_probable_title(text: str) -> bool:
    if re.match(r"^[一二三四五六七八九十]+、", text):
        return False
    if re.match(r"^（[一二三四五六七八九十]+）", text) or re.match(r"^\([一二三四五六七八九十]+\)", text):
        return False
    if re.match(r"^\d+[\.．、]", text):
        return False
    return not is_salutation(text)


def font_for_level(template: BuiltInTemplate, level: str) -> tuple[str, int]:
    if level == "title":
        return template.title_font, template.title_size
    if level == "heading1":
        return template.heading_font, template.heading_size
    if level == "heading2":
        return template.second_heading_font, template.heading_size
    if level == "heading3":
        return template.third_heading_font, template.heading_size
    return template.body_font, template.body_size


def apply_run_font(run, east_asia_font: str, latin_font: str, size: int) -> None:
    run.style = None
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.font.bold = False
    run.font.italic = False
    run.font.underline = False
    run.font.highlight_color = None
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:highlight", "w:shd", "w:color"):
        for element in list(r_pr.findall(qn(tag))):
            r_pr.remove(element)
    set_bool_run_property(r_pr, "w:b", False)
    set_bool_run_property(r_pr, "w:bCs", False)
    set_bool_run_property(r_pr, "w:i", False)
    set_bool_run_property(r_pr, "w:iCs", False)
    set_underline_none(r_pr)
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def set_first_line_indent_chars(paragraph, chars: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    ind = p_pr.ind
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    for attr in ("w:firstLine", "w:hanging"):
        if ind.get(qn(attr)) is not None:
            del ind.attrib[qn(attr)]
    if chars > 0:
        ind.set(qn("w:firstLineChars"), str(chars * 100))
    elif ind.get(qn("w:firstLineChars")) is not None:
        del ind.attrib[qn("w:firstLineChars")]


def set_paragraph_line_spacing(paragraph, before_lines: int, after_lines: int) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    spacing = p_pr.spacing
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:beforeLines"), str(before_lines * 100))
    spacing.set(qn("w:afterLines"), str(after_lines * 100))
    for attr in ("w:beforeAutospacing", "w:afterAutospacing"):
        if spacing.get(qn(attr)) is not None:
            del spacing.attrib[qn(attr)]


def set_bool_run_property(r_pr, tag: str, enabled: bool) -> None:
    elements = r_pr.findall(qn(tag))
    element = elements[0] if elements else OxmlElement(tag)
    for extra in elements[1:]:
        r_pr.remove(extra)
    element.set(qn("w:val"), "1" if enabled else "0")
    if not elements:
        r_pr.append(element)


def set_underline_none(r_pr) -> None:
    elements = r_pr.findall(qn("w:u"))
    element = elements[0] if elements else OxmlElement("w:u")
    for extra in elements[1:]:
        r_pr.remove(extra)
    element.set(qn("w:val"), "none")
    if not elements:
        r_pr.append(element)


def normalize_text(text: str, template: BuiltInTemplate) -> str:
    value = text
    if template.normalize_spacing:
        value = re.sub(r"[ \t\u3000]+", " ", value).strip()
    if template.normalize_parentheses:
        value = re.sub(r"\(([一二三四五六七八九十]+)\)", r"（\1）", value)
    return value
