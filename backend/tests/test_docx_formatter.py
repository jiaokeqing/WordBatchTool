from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.schemas import BuiltInTemplate
from app.services.docx_formatter import apply_builtin_template, preview_template


def make_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("标题")
    document.add_paragraph("")
    document.add_paragraph("一、一级标题")
    document.add_paragraph("   ")
    document.add_paragraph("（一）二级标题")
    document.add_paragraph("1. 三级标题")
    document.add_paragraph("正文内容 (一)  with   spaces")
    document.save(path)


def test_apply_builtin_template(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    make_docx(source)

    apply_builtin_template(
        source,
        output,
        BuiltInTemplate(body_font="仿宋_GB2312", body_size=16, heading_font="黑体", heading_size=16),
    )

    assert output.exists()
    document = Document(output)
    section = document.sections[0]
    assert round(section.top_margin.cm, 1) == 3.7
    assert round(section.bottom_margin.cm, 1) == 3.5
    assert round(section.left_margin.cm, 1) == 2.8
    assert round(section.right_margin.cm, 1) == 2.6
    assert len(document.paragraphs) == 5
    assert document.paragraphs[0].alignment == 1
    assert document.paragraphs[1].paragraph_format.first_line_indent.pt == 0
    assert document.paragraphs[2].paragraph_format.first_line_indent.pt == 0
    assert document.paragraphs[3].paragraph_format.first_line_indent.pt == 0
    assert document.paragraphs[1].runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
    assert document.paragraphs[2].runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "楷体_GB2312"
    assert document.paragraphs[4].runs[0].font.size.pt == 16
    assert document.paragraphs[4].paragraph_format.line_spacing.pt == 28
    assert document.paragraphs[4].paragraph_format.first_line_indent.pt == 32
    assert "(一)" not in document.paragraphs[4].text
    assert "（一）" in document.paragraphs[4].text


def test_preview_template(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    make_docx(source)

    preview = preview_template(source)

    assert preview.paragraphs == 7
    assert preview.tables == 0
    assert "Normal" in preview.styles


def test_salutation_flush_left_body_indented_and_blank_lines_removed(tmp_path: Path) -> None:
    source = tmp_path / "salutation.docx"
    output = tmp_path / "salutation.out.docx"
    document = Document()
    document.add_paragraph("各位老师：")
    document.add_paragraph("")
    document.add_paragraph("这是正文第一段。")
    document.add_paragraph("　　")
    document.add_paragraph("这是正文第二段。")
    document.save(source)

    apply_builtin_template(source, output, BuiltInTemplate())

    result = Document(output)
    assert [paragraph.text for paragraph in result.paragraphs] == ["各位老师：", "这是正文第一段。", "这是正文第二段。"]
    assert result.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert result.paragraphs[0].paragraph_format.first_line_indent.pt == 0
    assert result.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert result.paragraphs[1].paragraph_format.first_line_indent.pt == 32
    assert result.paragraphs[2].paragraph_format.first_line_indent.pt == 32
